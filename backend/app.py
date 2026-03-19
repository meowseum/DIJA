import logging
import uuid
from datetime import date, timedelta
import re
from logging.handlers import RotatingFileHandler
from pathlib import Path
from typing import Dict, List

import eel
from docxtpl import DocxTemplate

from .models import ClassRecord, HolidayRange, LessonOverride, PostponeRecord, _parse_bool, parse_date
from .schedule import (
    _find_next_available_weekly,
    apply_overrides,
    apply_postpones,
    calculate_progress,
    generate_weekly_schedule,
    holiday_set,
)
from .sku import build_sku, parse_sku
from .config import data_file, get_template_dir, get_output_dir, get_eps_template_path
from .storage import (
    CLASS_HEADERS,
    backup_file,
    load_classes,
    load_app_config,
    load_holidays,
    load_overrides,
    load_postpones,
    load_settings,
    load_stock_history,
    save_app_config,
    save_classes,
    save_holidays,
    save_overrides,
    save_postpones,
    save_settings,
    save_stock_snapshot,
    load_eps_records,
    save_eps_records,
    load_eps_audit,
    save_eps_audit,
    list_eps_dates as _list_eps_dates,
    get_eps_after_total,
    get_eps_after_items,
    has_eps_records_for_date,
)


def _setup_logger() -> logging.Logger:
    log_dir = Path(__file__).resolve().parent.parent / "data"
    log_dir.mkdir(parents=True, exist_ok=True)
    logger = logging.getLogger("dij_log")
    if logger.handlers:
        return logger
    logger.setLevel(logging.INFO)
    handler = RotatingFileHandler(
        log_dir / "app.log", maxBytes=1_000_000, backupCount=3, encoding="utf-8"
    )
    handler.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
    logger.addHandler(handler)
    return logger


log = _setup_logger()


def _class_progress(
    class_record: ClassRecord,
    holidays: List[HolidayRange],
    postpones: List[PostponeRecord],
    overrides: List,
) -> dict:
    class_postpones = [p for p in postpones if p.class_id == class_record.id]
    class_overrides = [o for o in overrides if o.class_id == class_record.id]
    return calculate_progress(
        class_record.start_date,
        class_record.weekday,
        class_record.lesson_total,
        holidays,
        class_postpones,
        class_overrides,
    )


def _class_payload(
    class_record: ClassRecord,
    holidays: List[HolidayRange],
    postpones: List[PostponeRecord],
    overrides: List,
) -> dict:
    progress = _class_progress(class_record, holidays, postpones, overrides)
    payload = class_record.to_dict()
    payload.update(progress)
    return payload


def _build_schedule_with_index(
    class_record: ClassRecord,
    holidays: List[HolidayRange],
    postpones: List[PostponeRecord],
    overrides: List,
) -> list:
    base_schedule = generate_weekly_schedule(
        parse_date(class_record.start_date) or date.today(),
        class_record.weekday,
        class_record.lesson_total,
        holidays,
    )
    schedule = apply_postpones(base_schedule, class_record.weekday, holidays, postpones)
    schedule = apply_overrides(schedule, holidays, overrides)
    schedule.sort()
    return [{"date": lesson.isoformat(), "index": idx + 1} for idx, lesson in enumerate(schedule)]


def _make_up_date_for(
    class_record: ClassRecord,
    holidays: List[HolidayRange],
    postpones: List[PostponeRecord],
    overrides: List,
    original_date: date,
) -> str:
    schedule = generate_weekly_schedule(
        parse_date(class_record.start_date) or original_date,
        class_record.weekday,
        class_record.lesson_total,
        holidays,
    )
    schedule = apply_postpones(schedule, class_record.weekday, holidays, postpones)
    schedule = apply_overrides(schedule, holidays, overrides)
    scheduled = set(schedule)
    if original_date in scheduled:
        scheduled.remove(original_date)
    holidays_set = holiday_set(holidays)

    candidate = original_date + timedelta(days=7)
    if candidate in scheduled or candidate in holidays_set:
        candidate = _find_next_available_weekly(candidate, class_record.weekday, scheduled, holidays_set)
    return candidate.isoformat()


@eel.expose
def load_state() -> dict:
    classes = load_classes()
    holidays = load_holidays()
    postpones = load_postpones()
    settings = load_settings()
    app_config = load_app_config()
    overrides = load_overrides()

    # Data integrity check
    class_ids = {c.id for c in classes}
    for p in postpones:
        if p.class_id not in class_ids:
            log.warning("Orphaned postpone %s references missing class %s", p.id, p.class_id)
    for o in overrides:
        if o.class_id not in class_ids:
            log.warning("Orphaned override %s references missing class %s", o.id, o.class_id)
    for c in classes:
        if c.start_date and not parse_date(c.start_date):
            log.warning("Class %s has invalid start_date: %s", c.id, c.start_date)
        if not (0 <= c.weekday <= 6):
            log.warning("Class %s has out-of-range weekday: %s", c.id, c.weekday)

    updated = False
    today = date.today()
    for record in classes:
        if record.level == "初級" and record.relay_teacher and record.relay_date:
            relay_date = parse_date(record.relay_date)
            if relay_date and relay_date <= today and record.teacher != record.relay_teacher:
                record.teacher = record.relay_teacher
                updated = True
                log.info("Relay teacher applied for class %s", record.id)
    if updated:
        save_classes(classes)
    return {
        "classes": [_class_payload(c, holidays, postpones, overrides) for c in classes],
        "holidays": [h.to_dict() for h in holidays],
        "postpones": [p.to_dict() for p in postpones],
        "settings": settings,
        "app_config": app_config,
        "stock_history": load_stock_history(),
    }


@eel.expose
def set_app_location(location: str) -> dict:
    allowed = {"", "K", "L", "H"}
    location_code = (location or "").strip().upper()
    if location_code not in allowed:
        return {"ok": False, "error": "地點不正確。"}
    config = load_app_config()
    config["location"] = location_code
    save_app_config(config)
    return {"ok": True, "location": location_code}


@eel.expose
def set_tab_order(order: List[str]) -> dict:
    if not isinstance(order, list):
        return {"ok": False, "error": "排序格式不正確。"}
    clean = [str(item).strip() for item in order if str(item).strip()]
    config = load_app_config()
    config["tab_order"] = ",".join(clean)
    save_app_config(config)
    return {"ok": True}


@eel.expose
def set_eps_output_path(path: str) -> dict:
    output_path = (path or "").strip()
    config = load_app_config()
    config["eps_output_path"] = output_path
    save_app_config(config)
    return {"ok": True, "eps_output_path": output_path}


@eel.expose
def create_class(data: Dict) -> dict:
    class_code = data.get("sku", "").strip()
    sku_parts = parse_sku(class_code)
    if not sku_parts:
        return {"ok": False, "error": "班別格式不正確。"}

    start_date = data.get("start_date", "").strip()
    if not parse_date(start_date):
        return {"ok": False, "error": "開課日期不正確。"}

    lesson_total = int(data.get("lesson_total", 0))
    if lesson_total <= 0:
        return {"ok": False, "error": "總課節必須大於 0。"}

    selected_level = data.get("level", "").strip()
    if not selected_level:
        return {"ok": False, "error": "請選擇等級。"}
    if sku_parts["level"] and selected_level != sku_parts["level"]:
        return {"ok": False, "error": "等級與班別不一致。"}

    app_location = load_app_config().get("location", "").strip().upper()
    location_code = sku_parts["location"] or (app_location if app_location in {"K", "L", "H"} else "")
    sku_value = build_sku(
        selected_level,
        location_code,
        sku_parts["start_month"],
        sku_parts["class_letter"],
        sku_parts["start_year"],
    )

    relay_teacher = data.get("relay_teacher", "").strip()
    relay_date = data.get("relay_date", "").strip()
    if selected_level != "初級":
        relay_teacher = ""
        relay_date = ""
    elif relay_teacher and not relay_date:
        return {"ok": False, "error": "設有接力老師時，接力時間不能為空。"}
    elif relay_date and not parse_date(relay_date):
        return {"ok": False, "error": "接力時間不正確。"}

    class_record = ClassRecord(
        id=str(uuid.uuid4()),
        sku=sku_value,
        level=selected_level,
        location=location_code,
        start_month=sku_parts["start_month"],
        class_letter=sku_parts["class_letter"],
        start_year=sku_parts["start_year"],
        classroom=data.get("classroom", "").strip(),
        start_date=start_date,
        weekday=int(data.get("weekday", 0)),
        start_time=data.get("start_time", "").strip(),
        teacher=data.get("teacher", "").strip(),
        relay_teacher=relay_teacher,
        relay_date=relay_date,
        student_count=int(data.get("student_count", 0)),
        lesson_total=lesson_total,
        status="active",
        doorplate_done=False,
        questionnaire_done=False,
        intro_done=False,
    )

    classes = load_classes()
    classes.append(class_record)
    save_classes(classes)
    log.info("Class created: %s (id=%s)", sku_value, class_record.id)
    return {"ok": True}


@eel.expose
def update_class(class_id: str, updates: Dict) -> dict:
    classes = load_classes()
    updated = False
    bool_fields = {"doorplate_done", "questionnaire_done", "intro_done"}
    for index, record in enumerate(classes):
        if record.id != class_id:
            continue
        if "sku" in updates:
            sku_value = str(updates.get("sku") or "").strip()
            sku_parts = parse_sku(sku_value)
            if not sku_parts:
                return {"ok": False, "error": "班別格式不正確。"}
            level_value = sku_parts["level"] or record.level
            location_value = sku_parts["location"]
            record.level = level_value
            record.location = location_value
            record.start_month = sku_parts["start_month"]
            record.class_letter = sku_parts["class_letter"]
            record.start_year = sku_parts["start_year"]
            record.sku = build_sku(
                level_value,
                location_value,
                sku_parts["start_month"],
                sku_parts["class_letter"],
                sku_parts["start_year"],
            )
            updates = {k: v for k, v in updates.items() if k != "sku"}
        for field, value in updates.items():
            if hasattr(record, field):
                if field in bool_fields:
                    setattr(record, field, _parse_bool(value))
                else:
                    setattr(record, field, value)
        classes[index] = record
        updated = True
        break
    if updated:
        save_classes(classes)
        log.info("Class updated: %s fields=%s", class_id, list(updates.keys()) if isinstance(updates, dict) else "")
    return {"ok": updated}


@eel.expose
def delete_class(class_id: str) -> dict:
    classes = load_classes()
    if not any(record.id == class_id for record in classes):
        return {"ok": False, "error": "找不到班別。"}
    classes = [record for record in classes if record.id != class_id]
    save_classes(classes)
    log.info("Class deleted: %s", class_id)

    postpones = load_postpones()
    new_postpones = [postpone for postpone in postpones if postpone.class_id != class_id]
    if len(new_postpones) != len(postpones):
        save_postpones(new_postpones)

    overrides = load_overrides()
    new_overrides = [override for override in overrides if override.class_id != class_id]
    if len(new_overrides) != len(overrides):
        save_overrides(new_overrides)

    return {"ok": True}


@eel.expose
def end_class_action(class_id: str, action: str, target_id: str = "", new_sku: str = "") -> dict:
    classes = load_classes()
    updated = False
    if action == "terminate":
        for record in classes:
            if record.id == class_id:
                record.status = "terminated"
                updated = True
                break
    elif action == "merge":
        if not any(c.id == target_id for c in classes):
            return {"ok": False, "error": "找不到合併目標班別。"}
        for record in classes:
            if record.id == class_id:
                record.status = "merged"
                record.merged_into_id = target_id
                updated = True
                break
    elif action == "promote":
        sku_parts = parse_sku(new_sku)
        if not sku_parts:
            return {"ok": False, "error": "升級班別格式不正確。"}
        if not sku_parts["level"]:
            return {"ok": False, "error": "升級班別需要包含等級。"}
        promoted_id = str(uuid.uuid4())
        base = next((c for c in classes if c.id == class_id), None)
        if not base:
            return {"ok": False, "error": "找不到班別。"}
        promoted = ClassRecord(
            id=promoted_id,
            sku=f"{sku_parts['level']}{sku_parts['code']}",
            level=sku_parts["level"],
            location=sku_parts["location"],
            start_month=sku_parts["start_month"],
            class_letter=sku_parts["class_letter"],
            start_year=sku_parts["start_year"],
            classroom=base.classroom,
            start_date=base.start_date,
            weekday=base.weekday,
            start_time=base.start_time,
            teacher=base.teacher,
            relay_teacher=base.relay_teacher,
            relay_date=base.relay_date,
            student_count=base.student_count,
            lesson_total=base.lesson_total,
            status="active",
            doorplate_done=base.doorplate_done,
            questionnaire_done=base.questionnaire_done,
            intro_done=base.intro_done,
        )
        for record in classes:
            if record.id == class_id:
                record.status = "promoted"
                record.promoted_to_id = promoted_id
                updated = True
                break
        if updated:
            classes.append(promoted)
    if updated:
        save_classes(classes)
    return {"ok": updated}


@eel.expose
def add_holiday(data: Dict) -> dict:
    holiday = HolidayRange(
        id=str(uuid.uuid4()),
        start_date=data.get("start_date", "").strip(),
        end_date=data.get("end_date", "").strip(),
        name=data.get("name", "").strip(),
    )
    holidays = load_holidays()
    holidays.append(holiday)
    save_holidays(holidays)
    return {"ok": True}


@eel.expose
def delete_holiday(holiday_id: str) -> dict:
    holidays = load_holidays()
    new_holidays = [holiday for holiday in holidays if holiday.id != holiday_id]
    save_holidays(new_holidays)
    return {"ok": True}


@eel.expose
def add_postpone(class_id: str, original_date: str, reason: str) -> dict:
    classes = load_classes()
    class_record = next((c for c in classes if c.id == class_id), None)
    if not class_record:
        return {"ok": False, "error": "找不到班別。"}

    original_dt = parse_date(original_date)
    if not original_dt:
        return {"ok": False, "error": "原定日期不正確。"}

    holidays = load_holidays()
    postpones = load_postpones()
    overrides = load_overrides()
    schedule = _build_schedule_with_index(class_record, holidays, postpones, overrides)
    schedule_dates = {item["date"] for item in schedule}
    if original_date not in schedule_dates:
        progress = calculate_progress(
            class_record.start_date,
            class_record.weekday,
            class_record.lesson_total,
            holidays,
            postpones,
            overrides,
        )
        is_ended = progress["lessons_remaining"] == 0 or class_record.status == "terminated"
        if not is_ended:
            return {"ok": False, "error": "原定日期不在日程。"}
        make_up_date = _make_up_date_for(class_record, holidays, postpones, overrides, original_dt)
        class_record.lesson_total += 1
        reactivated = class_record.status == "terminated"
        if reactivated:
            class_record.status = "active"
            log.info("Class %s reactivated via add_postpone", class_id)
        save_classes(classes)
    else:
        make_up_date = _make_up_date_for(class_record, holidays, postpones, overrides, original_dt)
        reactivated = False

    postpone = PostponeRecord(
        id=str(uuid.uuid4()),
        class_id=class_id,
        original_date=original_date,
        reason=reason.strip(),
        make_up_date=make_up_date,
    )
    postpones.append(postpone)
    save_postpones(postpones)
    log.info("Postpone added for class %s: %s → %s", class_id, original_date, make_up_date)
    return {"ok": True, "reactivated": reactivated}


@eel.expose
def add_postpone_manual(class_id: str, original_date: str, make_up_date: str, reason: str) -> dict:
    classes = load_classes()
    class_record = next((c for c in classes if c.id == class_id), None)
    if not class_record:
        return {"ok": False, "error": "找不到班別。"}
    if not parse_date(original_date):
        return {"ok": False, "error": "原定日期不正確。"}
    if not parse_date(make_up_date):
        return {"ok": False, "error": "補課日期不正確。"}

    holidays = load_holidays()
    postpones = load_postpones()
    overrides = load_overrides()
    schedule = _build_schedule_with_index(class_record, holidays, postpones, overrides)
    schedule_dates = {item["date"] for item in schedule}
    if original_date not in schedule_dates:
        progress = calculate_progress(
            class_record.start_date,
            class_record.weekday,
            class_record.lesson_total,
            holidays,
            postpones,
            overrides,
        )
        is_ended = progress["lessons_remaining"] == 0 or class_record.status == "terminated"
        if not is_ended:
            return {"ok": False, "error": "原定日期不在日程。"}
        if make_up_date in schedule_dates:
            return {"ok": False, "error": "補課日期重複。"}
        if parse_date(make_up_date) in holiday_set(holidays):
            return {"ok": False, "error": "補課日期遇到假期。"}
        class_record.lesson_total += 1
        reactivated = class_record.status == "terminated"
        if reactivated:
            class_record.status = "active"
            log.info("Class %s reactivated via add_postpone_manual", class_id)
        save_classes(classes)
    else:
        reactivated = False
    postpone = PostponeRecord(
        id=str(uuid.uuid4()),
        class_id=class_id,
        original_date=original_date,
        reason=reason.strip(),
        make_up_date=make_up_date,
    )
    postpones.append(postpone)
    save_postpones(postpones)
    log.info("Manual postpone added for class %s: %s → %s", class_id, original_date, make_up_date)
    return {"ok": True, "reactivated": reactivated}


@eel.expose
def get_make_up_date(class_id: str, original_date: str) -> dict:
    classes = load_classes()
    class_record = next((c for c in classes if c.id == class_id), None)
    if not class_record:
        return {"ok": False, "error": "找不到班別。"}
    original_dt = parse_date(original_date)
    if not original_dt:
        return {"ok": False, "error": "原定日期不正確。"}
    holidays = load_holidays()
    postpones = load_postpones()
    overrides = load_overrides()
    schedule = _build_schedule_with_index(class_record, holidays, postpones, overrides)
    schedule_dates = {item["date"] for item in schedule}
    if original_date not in schedule_dates:
        progress = calculate_progress(
            class_record.start_date,
            class_record.weekday,
            class_record.lesson_total,
            holidays,
            postpones,
            overrides,
        )
        is_ended = progress["lessons_remaining"] == 0 or class_record.status == "terminated"
        if not is_ended:
            return {"ok": False, "error": "原定日期不在日程。"}
    make_up_date = _make_up_date_for(class_record, holidays, postpones, overrides, original_dt)
    return {"ok": True, "make_up_date": make_up_date}


@eel.expose
def add_setting(entry_type: str, value: str) -> dict:
    entry_type = entry_type.strip().lower()
    value = value.strip()
    if entry_type not in {"teacher", "room", "level", "time"}:
        return {"ok": False, "error": "設定類型不正確。"}
    if not value:
        return {"ok": False, "error": "請輸入內容。"}
    settings = load_settings()
    if value not in settings.get(entry_type, []):
        settings[entry_type].append(value)
        save_settings(settings)
    return {"ok": True}


@eel.expose
def delete_setting(entry_type: str, value: str) -> dict:
    entry_type = entry_type.strip().lower()
    value = value.strip()
    settings = load_settings()
    if entry_type in settings and value in settings[entry_type]:
        settings[entry_type].remove(value)
        save_settings(settings)
    return {"ok": True}


@eel.expose
def move_setting(entry_type: str, value: str, direction: str) -> dict:
    entry_type = entry_type.strip().lower()
    value = value.strip()
    direction = direction.strip().lower()
    settings = load_settings()
    entries = settings.get(entry_type, [])
    if value not in entries:
        return {"ok": False, "error": "找不到設定。"}
    index = entries.index(value)
    if direction == "up" and index > 0:
        entries[index - 1], entries[index] = entries[index], entries[index - 1]
    elif direction == "down" and index < len(entries) - 1:
        entries[index + 1], entries[index] = entries[index], entries[index + 1]
    settings[entry_type] = entries
    save_settings(settings)
    return {"ok": True}


@eel.expose
def export_settings_csv() -> dict:
    settings = load_settings()
    lines = ["type,value"]
    for entry_type in ("teacher", "room", "level", "time"):
        for value in settings.get(entry_type, []):
            safe_value = value.replace('"', '""')
            lines.append(f'{entry_type},"{safe_value}"')
    level_prices = settings.get("level_price", {}) or {}
    for level, price in level_prices.items():
        safe_value = f"{level}|{price}".replace('"', '""')
        lines.append(f'level_price,"{safe_value}"')
    for name, price in (settings.get("textbook", {}) or {}).items():
        safe_value = f"{name}|{price}".replace('"', '""')
        lines.append(f'textbook,"{safe_value}"')
    for k, v in (settings.get("level_textbook", {}) or {}).items():
        if isinstance(v, list) and v:
            safe_value = f"{k}|{','.join(v)}".replace('"', '""')
            lines.append(f'level_textbook,"{safe_value}"')
    for k, v in (settings.get("level_next", {}) or {}).items():
        safe_value = f"{k}|{v}".replace('"', '""')
        lines.append(f'level_next,"{safe_value}"')
    for name, count in (settings.get("textbook_stock", {}) or {}).items():
        safe_value = f"{name}|{count}".replace('"', '""')
        lines.append(f'textbook_stock,"{safe_value}"')
    return {"ok": True, "content": "\n".join(lines)}


@eel.expose
def import_settings_csv(content: str) -> dict:
    import csv
    import io

    backup_file(data_file("settings.csv"))
    new_settings = {
        "teacher": [], "room": [], "level": [], "time": [],
        "level_price": {}, "message_category": {},
        "textbook": {}, "level_textbook": {}, "textbook_stock": {}, "level_next": {},
    }
    try:
        reader = csv.DictReader(io.StringIO(content))
        for row in reader:
            entry_type = (row.get("type") or "").strip().lower()
            value = (row.get("value") or "").strip()
            if entry_type == "level_price":
                parsed = value.split("|", 1) if "|" in value else value.split("=", 1)
                if len(parsed) == 2:
                    level = parsed[0].strip()
                    price = parsed[1].strip()
                    if level and price:
                        try:
                            new_settings["level_price"][level] = max(0, int(float(price)))
                        except ValueError:
                            pass
                continue
            if entry_type == "textbook":
                parts = value.split("|", 1)
                if len(parts) == 2:
                    name, price_str = parts[0].strip(), parts[1].strip()
                    if name:
                        try:
                            new_settings["textbook"][name] = max(0, int(float(price_str)))
                        except ValueError:
                            new_settings["textbook"][name] = 0
                continue
            if entry_type == "textbook_stock":
                parts = value.split("|", 1)
                if len(parts) == 2:
                    name, count_str = parts[0].strip(), parts[1].strip()
                    if name:
                        try:
                            new_settings["textbook_stock"][name] = max(0, int(float(count_str)))
                        except ValueError:
                            new_settings["textbook_stock"][name] = 0
                continue
            if entry_type == "level_textbook":
                parts = value.split("|", 1)
                if len(parts) == 2:
                    k = parts[0].strip()
                    tb_list = [n.strip() for n in parts[1].split(",") if n.strip()]
                    if k and tb_list:
                        new_settings["level_textbook"][k] = tb_list
                continue
            if entry_type == "level_next":
                parts = value.split("|", 1)
                if len(parts) == 2:
                    k, v = parts[0].strip(), parts[1].strip()
                    if k:
                        new_settings["level_next"][k] = v
                continue
            if entry_type in new_settings and isinstance(new_settings[entry_type], list) and value and value not in new_settings[entry_type]:
                new_settings[entry_type].append(value)
    except csv.Error:
        return {"ok": False, "error": "CSV 內容不正確。"}

    save_settings(new_settings)
    return {"ok": True}


@eel.expose
def set_level_price(level: str, price: int) -> dict:
    level = (level or "").strip()
    if not level:
        return {"ok": False, "error": "等級不能為空。"}
    try:
        price_value = max(0, int(float(price)))
    except (TypeError, ValueError):
        return {"ok": False, "error": "學費格式不正確。"}
    settings = load_settings()
    level_prices = settings.get("level_price", {}) or {}
    level_prices[level] = price_value
    settings["level_price"] = level_prices
    save_settings(settings)
    return {"ok": True}


@eel.expose
def adjust_level_prices(delta: int) -> dict:
    try:
        delta_value = int(delta)
    except (TypeError, ValueError):
        return {"ok": False, "error": "調整值不正確。"}
    settings = load_settings()
    level_prices = settings.get("level_price", {}) or {}
    for key in list(level_prices.keys()):
        try:
            level_prices[key] = max(0, int(level_prices[key]) + delta_value)
        except (TypeError, ValueError):
            level_prices[key] = max(0, delta_value)
    settings["level_price"] = level_prices
    save_settings(settings)
    return {"ok": True}


@eel.expose
def set_textbook(name: str, price: int) -> dict:
    name = (name or "").strip()
    if not name:
        return {"ok": False, "error": "教材名稱不能為空。"}
    try:
        price_value = max(0, int(float(price)))
    except (TypeError, ValueError):
        return {"ok": False, "error": "價格格式不正確。"}
    settings = load_settings()
    textbooks = settings.get("textbook", {}) or {}
    textbooks[name] = price_value
    settings["textbook"] = textbooks
    save_settings(settings)
    return {"ok": True}


@eel.expose
def delete_textbook(name: str) -> dict:
    name = (name or "").strip()
    settings = load_settings()
    changed = False
    for key in ("textbook", "textbook_stock"):
        d = settings.get(key, {}) or {}
        if name in d:
            del d[name]
            settings[key] = d
            changed = True
    # Remove textbook from level_textbook arrays
    lt = settings.get("level_textbook", {}) or {}
    lt_changed = False
    for lv in list(lt.keys()):
        v = lt[lv]
        if isinstance(v, list) and name in v:
            v.remove(name)
            lt_changed = True
            if not v:
                del lt[lv]
    settings["level_textbook"] = lt
    if changed or lt_changed:
        save_settings(settings)
    return {"ok": True}


@eel.expose
def set_textbook_stock(name: str, count: int) -> dict:
    name = (name or "").strip()
    if not name:
        return {"ok": False, "error": "教材名稱不能為空。"}
    try:
        count_value = max(0, int(float(count)))
    except (TypeError, ValueError):
        return {"ok": False, "error": "數量格式不正確。"}
    settings = load_settings()
    stock = settings.get("textbook_stock", {}) or {}
    stock[name] = count_value
    settings["textbook_stock"] = stock
    save_settings(settings)
    return {"ok": True}


@eel.expose
def save_monthly_stock(month: str, stock_data: dict) -> dict:
    """Save a monthly stock snapshot. month format: 'YYYY-MM'."""
    month = (month or "").strip()
    if not re.match(r"^\d{4}-\d{2}$", month):
        return {"ok": False, "error": "月份格式不正確（應為 YYYY-MM）。"}
    cleaned: dict = {}
    for name, count in (stock_data or {}).items():
        name = str(name).strip()
        if not name:
            continue
        try:
            cleaned[name] = max(0, int(float(count)))
        except (TypeError, ValueError):
            cleaned[name] = 0
    save_stock_snapshot(month, cleaned)
    return {"ok": True}


@eel.expose
def get_stock_history() -> dict:
    return load_stock_history()


@eel.expose
def set_level_textbook(level: str, textbook_names: list) -> dict:
    level = (level or "").strip()
    if not level:
        return {"ok": False, "error": "等級不能為空。"}
    names = [n.strip() for n in (textbook_names or []) if str(n).strip()]
    settings = load_settings()
    lt = settings.get("level_textbook", {}) or {}
    if names:
        lt[level] = names
    elif level in lt:
        del lt[level]
    settings["level_textbook"] = lt
    save_settings(settings)
    return {"ok": True}


@eel.expose
def set_level_next(level: str, next_level: str) -> dict:
    level = (level or "").strip()
    next_level = (next_level or "").strip()
    if not level:
        return {"ok": False, "error": "等級不能為空。"}
    settings = load_settings()
    ln = settings.get("level_next", {}) or {}
    if next_level:
        ln[level] = next_level
    elif level in ln:
        del ln[level]
    settings["level_next"] = ln
    save_settings(settings)
    return {"ok": True}


@eel.expose
def set_last_review_ts(ts: str) -> dict:
    config = load_app_config()
    config["last_review_ts"] = (ts or "").strip()
    save_app_config(config)
    return {"ok": True}


@eel.expose
def save_student_counts(updates: list) -> dict:
    """Bulk update student_count for multiple classes. updates = [{id, student_count}]"""
    classes = load_classes()
    class_map = {c.id: c for c in classes}
    for item in updates:
        class_id = (item.get("id") or "").strip()
        count = item.get("student_count", None)
        if class_id in class_map and count is not None:
            try:
                class_map[class_id].student_count = max(0, int(count))
            except (TypeError, ValueError):
                pass
    save_classes(list(class_map.values()))
    return {"ok": True}


@eel.expose
def load_payment_template() -> dict:
    template_dir = get_template_dir()
    candidates = [
        template_dir / "messages" / "payinstruc_wtsapp.txt",
        template_dir / "payinstruc_wtsapp.txt",
    ]
    for path in candidates:
        if path.exists():
            return {"ok": True, "content": path.read_text(encoding="utf-8")}
    docx_path = get_template_dir() / "payinstruc_wtsapp.docx"
    if docx_path.exists():
        try:
            from docx import Document
        except ImportError:
            return {"ok": False, "error": "未安裝 python-docx，無法讀取 docx。"}
        doc = Document(docx_path)
        content = "\n".join([para.text for para in doc.paragraphs])
        return {"ok": True, "content": content}
    return {"ok": False, "error": "找不到模板。"}


@eel.expose
def load_makeup_template() -> dict:
    template_dir = get_template_dir()
    candidates = [
        template_dir / "messages" / "補堂注意事項及安排.docx",
    ]
    for path in candidates:
        if path.exists():
            if path.suffix.lower() == ".txt":
                return {"ok": True, "content": path.read_text(encoding="utf-8")}
            try:
                from docx import Document
            except ImportError:
                return {"ok": False, "error": "未安裝 python-docx，無法讀取 docx。"}
            doc = Document(path)
            content = "\n".join([para.text for para in doc.paragraphs])
            return {"ok": True, "content": content}
    return {"ok": False, "error": "找不到模板。"}


LOCATION_NAME = {"K": "旺角校", "L": "太子校", "H": "香港校"}
WEEKDAYS_ZH = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
WEEKDAYS_JP = ["月曜日", "火曜日", "水曜日", "木曜日", "金曜日", "土曜日", "日曜日"]
WEEKDAYS_EN = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]


def _format_class_name(record: ClassRecord) -> str:
    name = record.sku
    parts = parse_sku(record.sku)
    if parts:
        year_short = str(parts["start_year"])[-2:]
        prefix = "" if record.level.startswith("N") else "N"
        name = (
            f"{prefix}{record.level}{parts['location']}"
            f"{parts['start_month']}{parts['class_letter']}/{year_short}"
        )
    return name


def _format_class_time(time_value: str) -> str:
    class_time = time_value or ""
    match = re.search(r"(\d{1,2}):?(\d{2})\s*[-~]\s*(\d{1,2}):?(\d{2})", class_time)
    if not match:
        return class_time

    def _label(hour_str: str, min_str: str) -> str:
        hour = int(hour_str)
        minute = int(min_str)
        is_pm = hour >= 12
        hour12 = ((hour + 11) % 12) + 1
        return f"{hour12}:{minute:02d}{'pm' if is_pm else 'am'}"

    return f"{_label(match.group(1), match.group(2))} ~ {_label(match.group(3), match.group(4))}"


def _format_class_time_zh(time_value: str) -> str:
    """Format time in Chinese style: 下午7:00～9:30"""
    class_time = time_value or ""
    match = re.search(r"(\d{1,2}):?(\d{2})\s*[-~]\s*(\d{1,2}):?(\d{2})", class_time)
    if not match:
        return class_time

    def _zh_label(hour_str: str, min_str: str) -> str:
        hour = int(hour_str)
        minute = int(min_str)
        prefix = "下午" if hour >= 12 else "上午"
        hour12 = ((hour + 11) % 12) + 1
        return prefix, f"{hour12}:{minute:02d}"

    prefix1, t1 = _zh_label(match.group(1), match.group(2))
    _, t2 = _zh_label(match.group(3), match.group(4))
    return f"{prefix1}{t1}～{t2}"


def _monthly_hours_from_time(time_value: str) -> int:
    """Estimate monthly hours: session duration in hours * 4 weeks."""
    class_time = time_value or ""
    match = re.search(r"(\d{1,2}):?(\d{2})\s*[-~]\s*(\d{1,2}):?(\d{2})", class_time)
    if not match:
        return 0
    start_h = int(match.group(1)) * 60 + int(match.group(2))
    end_h = int(match.group(3)) * 60 + int(match.group(4))
    duration_mins = end_h - start_h
    if duration_mins <= 0:
        return 0
    hours_per_lesson = duration_mins / 60
    return round(hours_per_lesson * 4)


def _teacher_for(record: ClassRecord, use_relay: bool) -> str:
    if use_relay and record.level == "初級" and record.relay_teacher:
        return record.relay_teacher
    return record.teacher or ""


@eel.expose
def list_docx_templates() -> dict:
    template_dir = get_template_dir() / "print"
    if not template_dir.exists():
        return {"ok": False, "error": "找不到模板資料夾。", "templates": []}
    templates = sorted(
        [
            path.name
            for path in template_dir.glob("*.docx")
            if path.is_file()
        ]
    )
    return {"ok": True, "templates": templates}


@eel.expose
def generate_docx(
    template_name: str,
    class_id: str,
    use_relay_teacher: bool = False,
    class_id_secondary: str = "",
    use_relay_teacher_secondary: bool = False,
) -> dict:
    template_dir = get_template_dir() / "print"
    output_dir = get_output_dir()
    template_path = template_dir / (template_name or "")
    if not template_path.exists():
        return {"ok": False, "error": "找不到模板。"}

    classes = load_classes()
    class_record = next((cls for cls in classes if cls.id == class_id), None)
    if not class_record:
        return {"ok": False, "error": "找不到班別。"}

    weekday_index = max(0, min(6, int(class_record.weekday)))

    context = {
        "CLASS_NAME": _format_class_name(class_record),
        "CLASS_WEEK": WEEKDAYS_ZH[weekday_index],
        "TEACHER_NAME": _teacher_for(class_record, use_relay_teacher),
        "CLASS_TIME": _format_class_time(class_record.start_time or ""),
        "ROOM_NUMBER": class_record.classroom or "",
        "WEEK_DAY": WEEKDAYS_JP[weekday_index],
    }

    if template_name == "class.docx":
        context = {
            "CLASS_NAME_1": context["CLASS_NAME"],
            "CLASS_WEEK_1": WEEKDAYS_EN[weekday_index],
            "TEACHER_NAME_1": context["TEACHER_NAME"],
            "CLASS_TIME_1": context["CLASS_TIME"],
            "CLASS_NAME_2": "",
            "CLASS_WEEK_2": "",
            "TEACHER_NAME_2": "",
            "CLASS_TIME_2": "",
        }
        if class_id_secondary:
            secondary = next((cls for cls in classes if cls.id == class_id_secondary), None)
            if not secondary:
                return {"ok": False, "error": "找不到追加班別。"}
            sec_weekday = max(0, min(6, int(secondary.weekday)))
            context.update(
                {
                    "CLASS_NAME_2": _format_class_name(secondary),
                    "CLASS_WEEK_2": WEEKDAYS_EN[sec_weekday],
                    "TEACHER_NAME_2": _teacher_for(secondary, use_relay_teacher_secondary),
                    "CLASS_TIME_2": _format_class_time(secondary.start_time or ""),
                }
            )

    output_dir.mkdir(parents=True, exist_ok=True)
    stamp = date.today().strftime("%Y%m%d")
    output_name = f"{class_record.sku}_{template_path.stem}_{stamp}.docx"
    output_path = output_dir / output_name

    doc = DocxTemplate(template_path)
    doc.render(context)
    doc.save(output_path)
    return {"ok": True, "path": str(output_path)}


@eel.expose
def open_output_folder() -> dict:
    import os
    output_dir = get_output_dir()
    try:
        os.startfile(str(output_dir))
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@eel.expose
def get_promote_notice_data(class_id: str) -> dict:
    """Return auto-filled promotion notice fields for the given (target) class."""
    classes = load_classes()
    record = next((c for c in classes if c.id == class_id), None)
    if not record:
        return {"ok": False, "error": "找不到班別。"}

    settings = load_settings()
    holidays = load_holidays()
    postpones = load_postpones()
    overrides = load_overrides()

    # Compute end date from schedule
    start = parse_date(record.start_date)
    end_date_str = ""
    if start:
        class_postpones = [p for p in postpones if p.class_id == record.id]
        class_overrides = [o for o in overrides if o.class_id == record.id]
        schedule = generate_weekly_schedule(start, record.weekday, record.lesson_total, holidays)
        schedule = apply_postpones(schedule, record.weekday, holidays, class_postpones)
        schedule = apply_overrides(schedule, holidays, class_overrides)
        if schedule:
            end = schedule[-1]
            end_date_str = (
                f"{end.year}年{end.month}月{end.day}日"
            )

    # Format start date
    start_date_formatted = ""
    if start:
        weekday_index = max(0, min(6, record.weekday))
        start_date_formatted = (
            f"{start.month}月{start.day}日{WEEKDAYS_ZH[weekday_index]}"
        )

    # Duration string
    duration = ""
    if start and end_date_str:
        duration = f"{start.year}年{start.month}月{start.day}日至{end_date_str}"

    # Time in Chinese format
    time_zh = _format_class_time_zh(record.start_time or "")

    # Teacher with brackets
    teacher = f"＜{record.teacher}＞" if record.teacher else ""

    # Location
    location = LOCATION_NAME.get(record.location, record.location)

    # Monthly hours
    hours = _monthly_hours_from_time(record.start_time or "")
    hours_str = f"每月{hours}小時" if hours else "每月--小時"

    # Price
    level_price = settings.get("level_price", {}) or {}
    price = level_price.get(record.level, "")
    price_str = f"學費${price}" if price else "學費$--"

    remarks = f"{hours_str} {price_str}"

    # Signature date
    signature_date = f"{record.start_month:02d}/{record.start_year}"

    # Source level (look up from promoted_from if available, else derive from level_next)
    level_next = settings.get("level_next", {}) or {}
    source_level = next((k for k, v in level_next.items() if v == record.level), "")

    # Short level label for body text (strip 級 suffix for compactness, e.g. 初級→初, 中級→中)
    def _short(lvl: str) -> str:
        return lvl.rstrip("級") if lvl else lvl

    return {
        "ok": True,
        "name": _format_class_name(record),
        "start_date_formatted": start_date_formatted,
        "duration": duration,
        "time": time_zh,
        "teacher": teacher,
        "location": location,
        "remarks": remarks,
        "signature_date": signature_date,
        "source_level": source_level,
        "target_level": record.level,
        "source_level_short": _short(source_level),
        "target_level_short": _short(record.level),
        "start_month": record.start_month,
    }


@eel.expose
def generate_promote_notice(data: dict) -> dict:
    """Render promote_notice.docx template and save to output dir."""
    template_path = get_template_dir() / "print" / "promote_notice.docx"
    if not template_path.exists():
        return {"ok": False, "error": "找不到升班通知模板 (promote_notice.docx)。請先生成模板。"}

    output_dir = get_output_dir()
    output_dir.mkdir(parents=True, exist_ok=True)

    sku = (data.get("sku") or "promote").replace("/", "-")
    stamp = date.today().strftime("%Y%m%d")
    output_name = f"{sku}_promote_notice_{stamp}.docx"
    output_path = output_dir / output_name

    context = {
        "ADDRESSEE": data.get("addressee", ""),
        "BODY_TEXT": data.get("body_text", ""),
        "CLASS_NAME": data.get("name", ""),
        "START_DATE": data.get("start_date_formatted", ""),
        "DURATION": data.get("duration", ""),
        "TIME": data.get("time", ""),
        "TEACHER": data.get("teacher", ""),
        "LOCATION": data.get("location", ""),
        "REMARKS": data.get("remarks", ""),
        "TEXTBOOK_FEE": data.get("textbook_fee", ""),
        "SIGNATURE_DATE": data.get("signature_date", ""),
    }

    try:
        doc = DocxTemplate(template_path)
        doc.render(context)
        doc.save(output_path)
        return {"ok": True, "path": str(output_path)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@eel.expose
def list_message_templates() -> dict:
    template_dir = get_template_dir() / "messages"
    if not template_dir.exists():
        return {"ok": False, "error": "找不到訊息資料夾。", "templates": []}
    settings = load_settings()
    category_map = settings.get("message_category", {}) or {}
    templates = []
    for path in sorted(template_dir.glob("*")):
        if not path.is_file():
            continue
        if path.suffix.lower() not in {".txt", ".docx"}:
            continue
        if path.name == "補堂注意事項及安排.docx":
            continue
        templates.append(
            {
                "name": path.name,
                "label": path.stem,
                "category": category_map.get(path.name, ""),
            }
        )
    return {"ok": True, "templates": templates}


@eel.expose
def load_message_content(template_name: str) -> dict:
    template_dir = get_template_dir() / "messages"
    template_path = template_dir / (template_name or "")
    if not template_path.exists():
        return {"ok": False, "error": "找不到訊息模板。"}
    if template_path.suffix.lower() == ".txt":
        return {"ok": True, "content": template_path.read_text(encoding="utf-8")}
    if template_path.suffix.lower() == ".docx":
        try:
            from docx import Document
        except ImportError:
            return {"ok": False, "error": "未安裝 python-docx，無法讀取 docx。"}
        doc = Document(template_path)
        content = "\n".join([para.text for para in doc.paragraphs])
        return {"ok": True, "content": content}
    return {"ok": False, "error": "不支援的檔案格式。"}


@eel.expose
def set_message_category(template_name: str, category: str) -> dict:
    name = (template_name or "").strip()
    if not name:
        return {"ok": False, "error": "模板名稱不正確。"}
    settings = load_settings()
    category_map = settings.get("message_category", {}) or {}
    category_map[name] = (category or "").strip()
    settings["message_category"] = category_map
    save_settings(settings)
    return {"ok": True}


@eel.expose
def delete_postpone(postpone_id: str) -> dict:
    postpones = load_postpones()
    new_postpones = [postpone for postpone in postpones if postpone.id != postpone_id]
    save_postpones(new_postpones)
    return {"ok": True}


@eel.expose
def get_class_schedule(class_id: str) -> dict:
    classes = load_classes()
    class_record = next((c for c in classes if c.id == class_id), None)
    if not class_record:
        return {"ok": False, "error": "找不到班別。"}
    holidays = load_holidays()
    postpones = load_postpones()
    overrides = load_overrides()
    class_postpones = [p for p in postpones if p.class_id == class_id]
    class_overrides = [o for o in overrides if o.class_id == class_id]
    schedule_payload = []
    make_up_dates = {p.make_up_date for p in class_postpones}
    schedule_with_index = _build_schedule_with_index(class_record, holidays, class_postpones, class_overrides)
    for item in schedule_with_index:
        date_str = item["date"]
        schedule_payload.append(
            {
                "date": date_str,
                "type": "makeup" if date_str in make_up_dates else "normal",
                "index": item["index"],
                "total": class_record.lesson_total,
            }
        )
    return {
        "ok": True,
        "class": class_record.to_dict(),
        "schedule": schedule_payload,
        "postpones": [p.to_dict() for p in class_postpones],
        "overrides": [o.to_dict() for o in class_overrides],
    }


@eel.expose
def add_schedule_override(class_id: str, date_str: str, action: str) -> dict:
    if not parse_date(date_str):
        return {"ok": False, "error": "日期不正確。"}
    if action not in {"add", "remove"}:
        return {"ok": False, "error": "操作不正確。"}
    overrides = load_overrides()
    exists = next((o for o in overrides if o.class_id == class_id and o.date == date_str and o.action == action), None)
    if exists:
        return {"ok": True}
    overrides.append(
        LessonOverride(
            id=str(uuid.uuid4()),
            class_id=class_id,
            date=date_str,
            action=action,
        )
    )
    save_overrides(overrides)
    return {"ok": True}


@eel.expose
def delete_schedule_override(override_id: str) -> dict:
    overrides = load_overrides()
    new_overrides = [o for o in overrides if o.id != override_id]
    save_overrides(new_overrides)
    return {"ok": True}


@eel.expose
def terminate_class_with_last_date(class_id: str, last_date: str) -> dict:
    end_date = parse_date(last_date)
    if not end_date:
        return {"ok": False, "error": "日期不正確。"}
    classes = load_classes()
    class_record = next((c for c in classes if c.id == class_id), None)
    if not class_record:
        return {"ok": False, "error": "找不到班別。"}
    holidays = load_holidays()
    postpones = load_postpones()
    overrides = load_overrides()
    schedule = _build_schedule_with_index(class_record, holidays, postpones, overrides)
    count = sum(1 for item in schedule if parse_date(item["date"]) and parse_date(item["date"]) <= end_date)
    if count <= 0:
        return {"ok": False, "error": "結束日期早於開課日期。"}
    class_record.lesson_total = count
    class_record.status = "terminated"
    save_classes(classes)
    return {"ok": True}


@eel.expose
def get_calendar_data(start_date: str, end_date: str) -> dict:
    start = parse_date(start_date)
    end = parse_date(end_date)
    if not start or not end:
        return {"ok": False, "error": "日期不正確。"}
    if end < start:
        start, end = end, start
    classes = load_classes()
    holidays = load_holidays()
    postpones = load_postpones()
    overrides = load_overrides()
    sessions = []
    for class_record in classes:
        class_postpones = [p for p in postpones if p.class_id == class_record.id]
        class_overrides = [o for o in overrides if o.class_id == class_record.id]
        schedule = _build_schedule_with_index(class_record, holidays, class_postpones, class_overrides)
        for item in schedule:
            lesson_date = parse_date(item["date"])
            if not lesson_date:
                continue
            if lesson_date < start or lesson_date > end:
                continue
            idx = item["index"]
            payment_due = idx % 4 in {3, 0}
            sessions.append(
                {
                    "date": item["date"],
                    "sku": class_record.sku,
                    "class_id": class_record.id,
                    "location": class_record.location,
                    "room": class_record.classroom,
                    "teacher": class_record.teacher,
                    "time": class_record.start_time,
                    "lesson_index": idx,
                    "lesson_total": class_record.lesson_total,
                    "payment_due": payment_due,
                }
            )
    return {
        "ok": True,
        "sessions": sessions,
        "holidays": [h.to_dict() for h in holidays],
    }


@eel.expose
def export_classes_csv() -> dict:
    import csv
    import io

    classes = load_classes()
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=CLASS_HEADERS)
    writer.writeheader()
    for record in classes:
        writer.writerow(record.to_dict())
    return {"ok": True, "content": output.getvalue()}


@eel.expose
def import_classes_csv(content: str) -> dict:
    import csv
    import io

    try:
        reader = csv.DictReader(io.StringIO(content))
        classes = [ClassRecord.from_dict(row) for row in reader]
    except csv.Error:
        return {"ok": False, "error": "CSV 內容不正確。"}
    backup_file(data_file("classes.csv"))
    save_classes(classes)
    log.info("Classes CSV imported: %d records", len(classes))
    return {"ok": True}


# ---------------------------------------------------------------------------
# EPS audit
# ---------------------------------------------------------------------------

_eps_items_cache: list | None = None


def _parse_hk_price(s: str) -> int | None:
    s = s.strip().replace("HK$", "").replace(",", "").replace(" ", "")
    if not s:
        return None
    try:
        return int(float(s))
    except ValueError:
        return None


def _parse_eps_template() -> list:
    """Parse the EPS blank template CSV and return list of item dicts."""
    global _eps_items_cache
    if _eps_items_cache is not None:
        return _eps_items_cache

    import csv as _csv
    path = get_eps_template_path()
    if not path.exists():
        log.warning("EPS template not found: %s", path)
        return []

    with path.open("r", encoding="utf-8") as f:
        rows = list(_csv.reader(f))

    items = []
    section = "class"
    for i, row in enumerate(rows):
        if i < 3:
            continue
        col_a = (row[0] if len(row) > 0 else "").strip()
        col_b = (row[1] if len(row) > 1 else "").strip()

        if col_a == "\u66f8":  # 書
            section = "book"
            continue
        if col_a == "\u5176\u4ed6":  # 其他
            section = "other"
            continue

        if not col_a or "Sub-total" in col_a or "Total" in col_a:
            continue
        # Check if any cell in the row contains Sub-total
        if any("Sub-total" in str(c) for c in row):
            continue

        price = _parse_hk_price(col_b)
        if price is None:
            continue

        items.append({
            "name": col_a,
            "price": price,
            "section": section,
        })

    _eps_items_cache = items
    return items


@eel.expose
def load_eps_items() -> dict:
    items = _parse_eps_template()
    return {"ok": True, "items": items}


@eel.expose
def load_eps_record(date_str: str) -> dict:
    items = _parse_eps_template()
    rows = load_eps_records(date_str)
    audit = load_eps_audit(date_str)

    # Build lookup: (item_name, period) -> row
    lookup = {}
    for r in rows:
        key = ((r.get("item_name") or "").strip(), (r.get("period") or "").strip())
        lookup[key] = r

    # Get yesterday's date for carry-over
    from datetime import datetime
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d").date()
        yesterday = (d - timedelta(days=1)).isoformat()
    except ValueError:
        yesterday = ""
    past_day_carry = get_eps_after_total(yesterday) if yesterday else 0

    # If no before-1900 records saved yet, carry over yesterday's after-1900 items
    carry_over = {}
    if yesterday and not has_eps_records_for_date(date_str, "before"):
        carry_over = get_eps_after_items(yesterday)

    records = {"before": [], "after": []}
    for item in items:
        for period in ("before", "after"):
            r = lookup.get((item["name"], period))
            qty_k = int(float(r.get("qty_K", "0"))) if r else 0
            qty_l = int(float(r.get("qty_L", "0"))) if r else 0
            qty_hk = int(float(r.get("qty_HK", "0"))) if r else 0

            # Merge carry-over into before-1900 (only when no saved records)
            if period == "before" and item["name"] in carry_over:
                co = carry_over[item["name"]]
                qty_k += co.get("qty_K", 0)
                qty_l += co.get("qty_L", 0)
                qty_hk += co.get("qty_HK", 0)

            records[period].append({
                "item_name": item["name"],
                "qty_K": qty_k,
                "qty_L": qty_l,
                "qty_HK": qty_hk,
            })

    def _audit_int(key):
        if not audit:
            return 0
        # Support both new (operator_1_before) and legacy (operator_1) field names
        val = audit.get(key, "0")
        try:
            return int(float(val))
        except (ValueError, TypeError):
            return 0

    audit_data = {
        "operator_1_before": _audit_int("operator_1_before") or _audit_int("operator_1"),
        "operator_2_before": _audit_int("operator_2_before") or _audit_int("operator_2"),
        "operator_3_before": _audit_int("operator_3_before") or _audit_int("operator_3"),
        "operator_1_after": _audit_int("operator_1_after"),
        "operator_2_after": _audit_int("operator_2_after"),
        "operator_3_after": _audit_int("operator_3_after"),
    }

    return {
        "ok": True,
        "records": records,
        "audit": audit_data,
        "past_day_carry": past_day_carry,
    }


@eel.expose
def save_eps_record(date_str: str, records: dict, audit: dict) -> dict:
    items = _parse_eps_template()

    csv_rows = []
    sheet_before = 0
    sheet_after = 0

    for period in ("before", "after"):
        period_items = records.get(period, [])
        for idx, entry in enumerate(period_items):
            if idx >= len(items):
                break
            item = items[idx]
            qty_k = max(0, int(entry.get("qty_K", 0)))
            qty_l = max(0, int(entry.get("qty_L", 0)))
            qty_hk = max(0, int(entry.get("qty_HK", 0)))
            subtotal = item["price"] * (qty_k + qty_l + qty_hk)

            if period == "before":
                sheet_before += subtotal
            else:
                sheet_after += subtotal

            csv_rows.append({
                "date": date_str,
                "item_name": item["name"],
                "item_price": str(item["price"]),
                "item_section": item["section"],
                "qty_K": str(qty_k),
                "qty_L": str(qty_l),
                "qty_HK": str(qty_hk),
                "subtotal": str(subtotal),
                "period": period,
            })

    save_eps_records(date_str, csv_rows)

    # Compute audit
    from datetime import datetime
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d").date()
        yesterday = (d - timedelta(days=1)).isoformat()
    except ValueError:
        yesterday = ""
    past_day_carry = get_eps_after_total(yesterday) if yesterday else 0

    op1b = max(0, int(audit.get("operator_1_before", 0)))
    op2b = max(0, int(audit.get("operator_2_before", 0)))
    op3b = max(0, int(audit.get("operator_3_before", 0)))
    op1a = max(0, int(audit.get("operator_1_after", 0)))
    op2a = max(0, int(audit.get("operator_2_after", 0)))
    op3a = max(0, int(audit.get("operator_3_after", 0)))
    ops_sum_before = op1b + op2b + op3b
    ops_sum_after = op1a + op2a + op3a
    calculated_total = ops_sum_before + past_day_carry
    status = "OK" if calculated_total == sheet_before else "MISMATCH"
    status_after = "OK" if ops_sum_after == sheet_after else "MISMATCH"
    status_audit = "OK" if (ops_sum_before + ops_sum_after) == (sheet_before + sheet_after) else "MISMATCH"

    audit_row = {
        "date": date_str,
        "operator_1_before": str(op1b),
        "operator_2_before": str(op2b),
        "operator_3_before": str(op3b),
        "operator_1_after": str(op1a),
        "operator_2_after": str(op2a),
        "operator_3_after": str(op3a),
        "operators_sum_before": str(ops_sum_before),
        "operators_sum_after": str(ops_sum_after),
        "sheet_before": str(sheet_before),
        "sheet_after": str(sheet_after),
        "past_day_carry": str(past_day_carry),
        "calculated_total": str(calculated_total),
        "status": status,
        "status_after": status_after,
        "status_audit": status_audit,
    }
    save_eps_audit(date_str, audit_row)

    log.info("EPS record saved for %s: status=%s status_after=%s status_audit=%s",
             date_str, status, status_after, status_audit)
    return {
        "ok": True,
        "status": status,
        "status_after": status_after,
        "status_audit": status_audit,
        "calculated_total": calculated_total,
        "operators_sum_before": ops_sum_before,
        "operators_sum_after": ops_sum_after,
        "past_day_carry": past_day_carry,
        "sheet_before": sheet_before,
        "sheet_after": sheet_after,
    }


@eel.expose
def export_eps_csv(date_str: str) -> dict:
    """Generate report-only HTML matching the Excel reference style."""
    from html import escape
    from datetime import datetime

    items = _parse_eps_template()
    rows = load_eps_records(date_str)

    # Merge before + after quantities per item
    merged: dict = {}
    for r in rows:
        name = (r.get("item_name") or "").strip()
        if name not in merged:
            merged[name] = {"qty_K": 0, "qty_L": 0, "qty_HK": 0}
        for loc in ("qty_K", "qty_L", "qty_HK"):
            try:
                merged[name][loc] += int(float(r.get(loc, "0")))
            except ValueError:
                pass

    day_names = ["\u661f\u671f\u4e00", "\u661f\u671f\u4e8c", "\u661f\u671f\u4e09",
                 "\u661f\u671f\u56db", "\u661f\u671f\u4e94", "\u661f\u671f\u516d", "\u661f\u671f\u65e5"]
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d").date()
        dow = day_names[d.weekday()]
        date_display = d.strftime("%Y/%m/%d")
    except ValueError:
        dow = ""
        date_display = date_str

    year = date_str[:4] if len(date_str) >= 4 else ""

    def hk_fmt(v: int) -> str:
        return f"HK${v:,}" if v else "HK$0"

    def price_fmt(v: int) -> str:
        return f"HK${v:,} " if v >= 1000 else f"HK${v} "

    # Build HTML
    h = []
    h.append("""<html><head><meta charset="utf-8">
<style>
@page { margin: 0.17in 0.16in 0.17in 0.12in; }
body { margin: 0; }
table { border-collapse: collapse; table-layout: fixed; }
td { font-family: "Times New Roman", serif; font-size: 14pt; vertical-align: middle;
     white-space: nowrap; padding: 1px 4px; }
.hdr-year { font-size: 12pt; text-align: right; }
.hdr-label { font-size: 12pt; text-align: left; }
.title { font-size: 16pt; font-weight: 700; text-align: center; }
.col-hdr { font-size: 16pt; font-weight: 700; text-align: center;
           border: 1pt solid black; border-bottom: .5pt solid black; }
.col-hdr-left { font-size: 16pt; border-top: 1pt solid black; border-bottom: .5pt solid black;
                border-left: 1pt solid black; }
.item-name { font-weight: 700; border-left: 1pt solid black; border-bottom: .5pt solid black;
             border-right: none; }
.item-price { font-weight: 700; text-align: left; border-bottom: .5pt solid black; }
.qty-x { text-align: right; border-left: .5pt solid black; border-bottom: .5pt solid black; }
.qty-n { text-align: center; border-bottom: .5pt solid black; }
.qty-eq { text-align: center; border-bottom: .5pt solid black; }
.qty-val { text-align: right; border-right: .5pt solid black; border-bottom: .5pt solid black; }
.row-total { text-align: right; border: .5pt solid black; }
.row-remark { text-align: left; border-right: 1pt solid black; border-bottom: .5pt solid black;
              border-left: .5pt solid black; }
.section-hdr { font-size: 16pt; font-weight: 700; text-align: left;
               border-top: 1pt solid black; border-left: 1pt solid black;
               border-bottom: .5pt solid black; }
.section-hdr-r { border-top: 1pt solid black; border-right: 1pt solid black;
                 border-bottom: .5pt solid black; }
.sub-label { font-weight: 700; text-align: right; }
.sub-val { text-align: right; border-bottom: .5pt solid black; }
.loc-total { font-size: 9pt; color: windowtext; text-align: left; }
.total-label { font-weight: 700; text-align: right; }
.total-val { text-align: right; border-top: .5pt solid black; border-bottom: 1pt solid black; }
.date-label { font-weight: 700; text-align: right; }
.date-val { font-weight: 700; text-align: right; border-bottom: 1pt solid black; }
.dow { font-weight: 700; text-align: center; }
.spacer-row td { height: 9pt; }
</style></head><body>
<table>
<col width=9><col width=241><col width=106>
<col width=21><col width=37><col width=21><col width=85>
<col width=21><col width=37><col width=21><col width=81>
<col width=21><col width=37><col width=21><col width=83>
<col width=94><col width=93>
""")

    # Row 1: Year + Date
    h.append(f"""<tr height=24>
 <td></td><td class="hdr-year">{escape(year)}</td><td class="hdr-label">\u5b78\u5e74</td>
 <td colspan=11></td>
 <td class="date-label">Date:</td><td class="date-val">{escape(date_display)}</td>
 <td class="dow">{escape(dow)}</td>
</tr>""")

    # Row 2: Title
    h.append(f"""<tr height=22>
 <td></td><td colspan=16 class="title">EPS \u6536\u652f\u7d00\u9304 (\u65fa\u89d2\u6821 - \u661f\u671f\u4e00\u81f3\u661f\u671f\u516d)</td>
</tr>""")

    # Row 3: Column headers
    h.append("""<tr height=28>
 <td></td><td class="col-hdr-left">&nbsp;</td><td class="col-hdr-left">&nbsp;</td>
 <td colspan=4 class="col-hdr">K</td>
 <td colspan=4 class="col-hdr" style="border-left:none">L</td>
 <td colspan=4 class="col-hdr" style="border-left:none">HK</td>
 <td class="col-hdr" style="border-left:none">Total</td>
 <td class="col-hdr" style="border-left:none">Remarks</td>
</tr>""")

    section_totals: dict = {"class": 0, "book": 0, "other": 0}
    loc_totals: dict = {"class": {"K": 0, "L": 0, "HK": 0},
                        "book": {"K": 0, "L": 0, "HK": 0},
                        "other": {"K": 0, "L": 0, "HK": 0}}
    current_section = None
    grand_total = 0

    def _subtotal_row(section: str) -> str:
        st = section_totals[section]
        lk = hk_fmt(loc_totals[section]["K"])
        ll = hk_fmt(loc_totals[section]["L"])
        lh = hk_fmt(loc_totals[section]["HK"])
        return f"""<tr height=26>
 <td></td><td></td><td></td>
 <td colspan=3 class="loc-total">{escape(lk)}</td><td></td>
 <td colspan=3 class="loc-total">{escape(ll)}</td><td></td>
 <td colspan=3 class="loc-total">{escape(lh)}</td>
 <td class="sub-label">Sub-total:</td><td class="sub-val">{st}</td><td></td>
</tr>"""

    def _spacer_row() -> str:
        return '<tr class="spacer-row"><td colspan=17></td></tr>'

    for item in items:
        if item["section"] != current_section:
            if current_section == "class":
                h.append(_subtotal_row("class"))
                h.append(_spacer_row())
                h.append(f"""<tr height=29>
 <td></td><td colspan=16 class="section-hdr">\u66f8</td>
</tr>""")
            elif current_section == "book" and item["section"] == "other":
                h.append(_subtotal_row("book"))
                h.append(_spacer_row())
                h.append(f"""<tr height=29>
 <td></td><td colspan=2 class="section-hdr">\u5176\u4ed6</td>
 <td colspan=4 class="section-hdr-r">&nbsp;</td>
 <td colspan=4 class="section-hdr-r">&nbsp;</td>
 <td colspan=4 class="section-hdr-r">&nbsp;</td>
 <td class="section-hdr-r">&nbsp;</td><td class="section-hdr-r">&nbsp;</td>
</tr>""")
            current_section = item["section"]

        m = merged.get(item["name"], {"qty_K": 0, "qty_L": 0, "qty_HK": 0})
        price = item["price"]
        sub_k = price * m["qty_K"]
        sub_l = price * m["qty_L"]
        sub_hk = price * m["qty_HK"]
        total = sub_k + sub_l + sub_hk
        section_totals[item["section"]] += total
        loc_totals[item["section"]]["K"] += sub_k
        loc_totals[item["section"]]["L"] += sub_l
        loc_totals[item["section"]]["HK"] += sub_hk
        grand_total += total

        qk_s = str(m["qty_K"]) if m["qty_K"] else "&nbsp;"
        ql_s = str(m["qty_L"]) if m["qty_L"] else "&nbsp;"
        qh_s = str(m["qty_HK"]) if m["qty_HK"] else "&nbsp;"

        h.append(f"""<tr height=22>
 <td></td>
 <td class="item-name">{escape(item["name"])}</td>
 <td class="item-price">{escape(price_fmt(price))}</td>
 <td class="qty-x">X</td><td class="qty-n">{qk_s}</td><td class="qty-eq">=</td><td class="qty-val">{sub_k}</td>
 <td class="qty-x">X</td><td class="qty-n">{ql_s}</td><td class="qty-eq">=</td><td class="qty-val">{sub_l}</td>
 <td class="qty-x">X</td><td class="qty-n">{qh_s}</td><td class="qty-eq">=</td><td class="qty-val">{sub_hk}</td>
 <td class="row-total">{total}</td>
 <td class="row-remark">&nbsp;</td>
</tr>""")

    # Final section sub-total
    if current_section:
        h.append(_subtotal_row(current_section))

    # Grand total row
    h.append(f"""<tr height=21>
 <td></td><td></td><td></td>
 <td colspan=11></td>
 <td colspan=2 class="total-label">Total:</td>
 <td class="total-val">{grand_total}</td><td></td>
</tr>""")

    h.append("</table></body></html>")

    filename = f"EPS_{date_str}.htm"
    content = "\n".join(h)

    # Save to custom output path if configured
    config = load_app_config()
    eps_output_path = config.get("eps_output_path", "").strip()
    saved_path = ""
    if eps_output_path:
        out_dir = Path(eps_output_path)
        if out_dir.is_dir():
            out_file = out_dir / filename
            out_file.write_text(content, encoding="utf-8")
            saved_path = str(out_file)

    return {"ok": True, "content": content, "filename": filename, "saved_path": saved_path}


@eel.expose
def list_eps_dates_endpoint() -> dict:
    return {"ok": True, "dates": _list_eps_dates()}

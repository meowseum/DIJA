"""
Run once to generate template/print/promote_notice.docx
Usage: python create_promote_template.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(__file__).parent / "template" / "print" / "promote_notice.docx"


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell. kwargs: top, bottom, left, right with dict {val, sz, color}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            cfg = kwargs[edge]
            tag.set(qn("w:val"), cfg.get("val", "single"))
            tag.set(qn("w:sz"), str(cfg.get("sz", 4)))
            tag.set(qn("w:color"), cfg.get("color", "000000"))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_run(para, text, bold=False, size=11, color=None, font_name="PMingLiU"):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # --- Greeting line ---
    p_greet = doc.add_paragraph()
    p_greet.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p_greet, "致{{ADDRESSEE}}，", bold=False, size=12)

    doc.add_paragraph()  # blank line

    # --- Body text ---
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p_body, "{{BODY_TEXT}}", size=12)

    doc.add_paragraph()  # blank line

    # --- Table ---
    headers = ["名稱", "開課日期", "全課程修讀期", "時間", "導師", "上課地點", "備註"]
    placeholders = [
        "{{CLASS_NAME}}",
        "{{START_DATE}}",
        "{{DURATION}}",
        "{{TIME}}",
        "{{TEACHER}}",
        "{{LOCATION}}",
        "{{REMARKS}}",
    ]

    table = doc.add_table(rows=2, cols=7)
    table.style = "Table Grid"

    border_cfg = {"val": "single", "sz": 4, "color": "000000"}

    # Header row
    header_bg = "B8CCE4"  # light blue matching reference
    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, top=border_cfg, bottom=border_cfg, left=border_cfg, right=border_cfg)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(para, hdr, bold=True, size=10, color="1F3864")

    # Data row
    data_bg = "E2EFDA"  # light purple/lavender matching reference
    for i, ph in enumerate(placeholders):
        cell = table.rows[1].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, data_bg)
        set_cell_border(cell, top=border_cfg, bottom=border_cfg, left=border_cfg, right=border_cfg)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(para, ph, size=10)

    # --- Textbook fee row (below main data, indented) ---
    p_tb = doc.add_paragraph()
    p_tb.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p_tb, "書本費：", bold=True, size=10)
    add_run(p_tb, "{{TEXTBOOK_FEE}}", size=10)

    doc.add_paragraph()  # blank line

    # --- Footer notice ---
    footer_text = (
        "同學收到成績後，如欲報讀升級課程，請即日辦理報名手續，未有即日報名的同學，將列入候補名單內，"
        "如需留位，必須於當日到校務處登記，另作安排。開課日不辦理新登記手續，由於學位所限，請同學合作。"
    )
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_footer = add_run(p_footer, footer_text, size=10)
    run_footer.underline = True

    doc.add_paragraph()  # blank line
    doc.add_paragraph()

    # --- Signature ---
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_sig, "校務處", size=11)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_date, "{{SIGNATURE_DATE}}", size=11)

    doc.save(OUTPUT)
    print(f"Template saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
